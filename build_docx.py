import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

doc = docx.Document()

# Set Margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Colors
COLOR_NAVY = RGBColor(26, 54, 93)      # #1A365D Primary
COLOR_GOLD = RGBColor(212, 168, 67)    # #D4A843 Accent
COLOR_DARK = RGBColor(45, 55, 72)      # #2D3748 Body text

HEX_NAVY = '1A365D'
HEX_ALT_ROW = 'F7FAFC'

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = title_p.add_run('INFORME ESTRATÉGICO DE PRECIOS Y PUBLICIDAD')
run_title.font.name = 'Arial'
run_title.font.size = Pt(20)
run_title.font.bold = True
run_title.font.color.rgb = COLOR_NAVY

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = subtitle_p.add_run('Análisis Tributario 100% Fiel a Factura N° 60861 (Zhang Hermanos Ltda)\nCliente: América Pulido SPA')
run_sub.font.name = 'Arial'
run_sub.font.size = Pt(11)
run_sub.font.color.rgb = COLOR_GOLD
run_sub.font.bold = True

doc.add_paragraph() # Spacer

# Section 1
h1 = doc.add_heading(level=1)
r_h1 = h1.add_run('1. Explicación Tributaria de la Factura (En palabras simples)')
r_h1.font.name = 'Arial'
r_h1.font.color.rgb = COLOR_NAVY
r_h1.font.bold = True

p1 = doc.add_paragraph()
r1 = p1.add_run(
    'Al revisar la Factura Electrónica N° 60861 emitida por Zhang Hermanos Ltda (RUT 77.759.660-8), '
    'se verifican exactamente 15 líneas de productos listados:\n\n'
    '• Columna PRECIO en Factura = VALORES NETOS (Sin IVA).\n'
    '• Al sumar la columna TOTAL de la factura da exactamente $181.345 CLP (Neto).\n'
    '• Se le adiciona el 19% de IVA ($34.455 CLP), obteniendo el TOTAL cobrado en débito de $215.800 CLP.\n\n'
    '👉 Aclaración de Nombres en Factura vs. Nombre Comercial en Web:\n'
    '1. Piedra Cónica: Aparece literalmente en la línea 8 de la factura como "PIEDRA CONICA P/PULIR CARBURO SILICIO M14*2 #36" (Código PCM14236) con un precio Neto exacto de $4.990 ($5.938 con IVA).\n'
    '2. Base Soporte: Aparece en la línea 10 como "DISCO SOPORTE 4" (Código 106378) a $4.900 Neto ($5.831 con IVA).\n'
    '3. Trompo de Resina: No se emite con la palabra "Trompo" en la factura del proveedor; el proveedor lo factura dentro de sus códigos de pulido (como "DISCO PULIR 4.5 KAMASA GRIS" a $9.180 Neto).'
)
r1.font.name = 'Arial'
r1.font.size = Pt(10.5)
r1.font.color.rgb = COLOR_DARK

doc.add_paragraph()

# Section 2 Header
h2 = doc.add_heading(level=1)
r_h2 = h2.add_run('2. Tabla de Auditoría: 15 Ítems de la Factura N° 60861 vs. Precios Web')
r_h2.font.name = 'Arial'
r_h2.font.color.rgb = COLOR_NAVY
r_h2.font.bold = True

p_desc = doc.add_paragraph()
r_desc = p_desc.add_run(
    'La siguiente tabla audita literalmente las 15 líneas impresas en la Factura N° 60861 con su precio Neto exacto, '
    'el 19% de IVA real pagado, el costo total bruto de bodega y el precio comercial publicado en la web.'
)
r_desc.font.name = 'Arial'
r_desc.font.size = Pt(10)
r_desc.font.color.rgb = COLOR_DARK

# Table Data: Exact 15 items on Factura N° 60861
headers = [
    'Código Factura', 
    'Detalle Literal en Factura', 
    'Precio Neto (Factura)', 
    '+ 19% IVA', 
    'Costo Real (con IVA)', 
    'Precio Web Oficial', 
    'Estado en Factura'
]

data = [
    ['106376', 'DISCO PULIR DIAMANTADO 4', '$1.870', '$355', '$2.225', '$4.990', 'Línea 1 en Factura N° 60861'],
    ['DPD115SC', 'DISCO PULIR DIAMANTADA SECO (Resina Seco)', '$3.360', '$638', '$3.998', '$5.800', 'Línea 2 en Factura N° 60861'],
    ['101073', 'DISCO CERAMICA CONTINUO 4.5', '$2.250', '$428', '$2.678', '$5.990', 'Línea 3 en Factura N° 60861'],
    ['101175', 'DISCO CERAMICA KAMASA 4.5', '$2.050', '$390', '$2.440', '$4.990', 'Línea 4 en Factura N° 60861'],
    ['KM252', 'DISCO CERAMICA 7 1/4 KAMASA', '$5.930', '$1.127', '$7.057', '$12.990', 'Línea 5 en Factura N° 60861'],
    ['101074', 'DISCO CERAMICA CONTINUO 7 1/4 KAMASA', '$6.450', '$1.226', '$7.676', '$13.990', 'Línea 6 en Factura N° 60861'],
    ['DPG36', 'DISCO PULIR GRANO 36', '$2.250', '$428', '$2.678', '$5.490', 'Línea 7 en Factura N° 60861'],
    ['PCM14236', 'PIEDRA CONICA P/PULIR CARBURO SILICIO M14*2 #36', '$4.990', '$948', '$5.938', '$12.990', 'Línea 8 en Factura N° 60861'],
    ['DPAC115', 'DISCO P/PULIR ACERO INOXIDABLE 115MM', '$4.000', '$760', '$4.760', '$8.990', 'Línea 9 en Factura N° 60861'],
    ['106378', 'DISCO SOPORTE 4 (Base para pulir)', '$4.900', '$931', '$5.831', '$3.000 / $9.990', 'Línea 10 en Factura N° 60861'],
    ['101328', 'DISCO C/LIJA 4.5', '$1.530', '$291', '$1.821', '$3.800', 'Línea 11 en Factura N° 60861'],
    ['DTG115', 'DISCO TRALAPADO 4.5" P/MARMOL PIEDRA', '$750', '$143', '$893', '$2.490', 'Línea 12 en Factura N° 60861'],
    ['104295', 'DISCO PULIR 4.5 KAMASA GRIS (Trompo / Abrasivo)', '$9.180', '$1.744', '$10.924', '$9.990', 'Línea 13 en Factura N° 60861'],
    ['SCDM35', 'SIERRA COPA DIAMANTADA M14X2, 35MM (Saca Bocado)', '$7.750', '$1.473', '$9.223', '$17.000', 'Línea 14 en Factura N° 60861'],
    ['DR11410', 'DISCO 114MM 10MM, P/RANURA (Ranurador)', '$32.000', '$6.080', '$38.080', '$59.990', 'Línea 15 en Factura N° 60861']
]

table = doc.add_table(rows=len(data) + 1, cols=len(headers))
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header Row
hdr_cells = table.rows[0].cells
for idx, title in enumerate(headers):
    hdr_cells[idx].text = title
    set_cell_background(hdr_cells[idx], HEX_NAVY)
    set_cell_margins(hdr_cells[idx], top=100, bottom=100, left=80, right=80)
    for p in hdr_cells[idx].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(8)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

# Data Rows
for row_idx, row_data in enumerate(data):
    row_cells = table.rows[row_idx + 1].cells
    bg_hex = HEX_ALT_ROW if row_idx % 2 == 1 else 'FFFFFF'
    for col_idx, cell_value in enumerate(row_data):
        row_cells[col_idx].text = cell_value
        set_cell_background(row_cells[col_idx], bg_hex)
        set_cell_margins(row_cells[col_idx], top=70, bottom=70, left=60, right=60)
        for p in row_cells[col_idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx in [0, 1] else WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(7.5)
                r.font.color.rgb = COLOR_DARK
                if col_idx in [2, 4, 5]:
                    r.font.bold = True

doc.add_paragraph()

# Section 3
h3 = doc.add_heading(level=1)
r_h3 = h3.add_run('3. Conclusión de Auditoría')
r_h3.font.name = 'Arial'
r_h3.font.color.rgb = COLOR_NAVY
r_h3.font.bold = True

p3 = doc.add_paragraph()
r3 = p3.add_run(
    '1. Piedra Cónica: SÍ está en la factura. Es el ítem 8 (`PCM14236`), descrito como "PIEDRA CONICA P/PULIR CARBURO SILICIO M14*2 #36" a $4.990 Neto ($5.938 con IVA).\n'
    '2. Trompo de Resina: No aparece con ese nombre comercial en la factura del proveedor; el proveedor lo factura bajo descripciones genéricas como `DISCO PULIR 4.5 KAMASA GRIS` ($9.180 Neto / $10.924 con IVA).\n'
    '3. Base Soporte: Aparece en el ítem 10 como `DISCO SOPORTE 4` (`106378`) a $4.900 Neto ($5.831 con IVA).'
)
r3.font.size = Pt(10.5)
r3.font.color.rgb = COLOR_DARK

try:
    doc.save('Informe_Estrategico_Precios_y_Publicidad_America_Pulido.docx')
    print('Saved to original name.')
except Exception as e:
    doc.save('Informe_Estrategico_Precios_y_Publicidad_America_Pulido_v2.docx')
    print('Saved to v2 name due to file lock.')
