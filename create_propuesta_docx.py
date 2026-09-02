import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

doc = docx.Document()

# Page Margins
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

# Brand Color Palette
COLOR_NAVY = RGBColor(10, 25, 47)       # #0A192F Dark Primary
COLOR_GOLD = RGBColor(212, 168, 67)     # #D4A843 Gold Brand Accent
COLOR_DARK = RGBColor(45, 55, 72)       # #2D3748 Body text
COLOR_GREEN = RGBColor(34, 197, 94)     # Profit Green

HEX_NAVY = '0A192F'
HEX_ALT_ROW = 'F8F9FA'

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=80, bottom=80, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# Title Header
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = title_p.add_run('INFORME COMPLETO DE AUDITORÍA, ESTUDIO DE MERCADO\nY MATRIZ DE PRECIOS RECOMENDADOS')
run_title.font.name = 'Arial'
run_title.font.size = Pt(18)
run_title.font.bold = True
run_title.font.color.rgb = COLOR_NAVY

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = sub_p.add_run('Auditoría Factura N° 60861 vs. Precios Web, benchmark Mercado Libre / Sodimac y Proyección Meta Ads\nCliente: América Pulido SPA')
run_sub.font.name = 'Arial'
run_sub.font.size = Pt(11)
run_sub.font.color.rgb = COLOR_GOLD
run_sub.font.bold = True

doc.add_paragraph()

# Section 1: Verification of Names & Factura Matching
h1 = doc.add_heading(level=1)
r_h1 = h1.add_run('1. Verificación de Coincidencia de Productos (Factura N° 60861 vs. Web)')
r_h1.font.name = 'Arial'
r_h1.font.color.rgb = COLOR_NAVY
r_h1.font.bold = True

p1 = doc.add_paragraph()
r1 = p1.add_run(
    'Al auditar la Factura Electrónica N° 60861 de Zhang Hermanos Ltda (Total $215.800 CLP con IVA), se observa una discrepancia '
    'entre los nombres genéricos del sistema de facturación mayorista y las descripciones comerciales de la web.\n\n'
    '👉 ¿Cómo identificamos exactamente cada producto?\n'
    '1. Código de Producto del Proveedor: Cada ítem tiene un código único asignado en la factura (ej. PCM14236, DPD115SC, SCDM35).\n'
    '2. Especificaciones Técnicas y Medidas: La factura indica diámetros y materiales exactos (ej. M14x2, 35mm, 114mm, Grano 36).\n'
    '3. Correlación de Fotografías de Bodega vs. Nombre Comercial:\n'
    '   - El "Trompo de resina" fue facturado bajo el código `104295` como `DISCO PULIR 4.5 KAMASA GRIS` a $9.180 Neto ($10.924 con IVA).\n'
    '   - La "Piedra cónica de desbaste" figura literalmente en la línea 8 como `PCM14236` `PIEDRA CONICA P/PULIR CARBURO SILICIO M14*2 #36` a $4.990 Neto ($5.938 con IVA).\n'
    '   - La "Base soporte" figura en la línea 10 como `106378` `DISCO SOPORTE 4` a $4.900 Neto ($5.831 con IVA).'
)
r1.font.name = 'Arial'
r1.font.size = Pt(10)
r1.font.color.rgb = COLOR_DARK

doc.add_paragraph()

# Section 2: Market Study & Benchmark Table
h2 = doc.add_heading(level=1)
r_h2 = h2.add_run('2. Matriz Comparativa Completa: Costo Bodega, Publicidad, Mercado y Precio Sugerido')
r_h2.font.name = 'Arial'
r_h2.font.color.rgb = COLOR_NAVY
r_h2.font.bold = True

p2 = doc.add_paragraph()
r2 = p2.add_run(
    'A continuación se presenta el estudio de mercado realizado en plataformas como Mercado Libre Chile, Sodimac y marmolerías especializadas. '
    'Se incorpora el Costo de Publicidad Meta Ads Estimado (CAC) por venta unitaria para garantizar que las ventas por internet dejen un retorno neto positivo.'
)
r2.font.name = 'Arial'
r2.font.size = Pt(9.5)
r2.font.color.rgb = COLOR_DARK

headers_comp = [
    'Código & Producto',
    'Costo Real (c/IVA)',
    'Meta Ads (CAC)',
    'Precio Web Actual',
    'Mercado Referencia',
    'Precio Sugerido',
    'Ganancia Neta Actual',
    'Ganancia Neta Sugerida'
]

data_comp = [
    ['106376 - Disco Pulir Diamantado 4"', '$2.225', '$1.200', '$4.990', '$4.500 - $6.000', '$4.990', '+$1.565', '+$1.565'],
    ['DPD115SC - Disco Resina Seco 4.5"', '$3.998', '$1.200', '$5.800', '$5.500 - $7.500', '$6.490', '+$602', '+$1.292 (+114%)'],
    ['101073 - Disco Continuo 4.5" Mármol', '$2.678', '$1.000', '$5.990', '$5.000 - $6.500', '$5.490', '+$2.312', '+$1.812'],
    ['101175 - Disco Segmentado Kamasa 4.5"', '$2.440', '$1.000', '$4.990', '$4.500 - $5.990', '$4.990', '+$1.550', '+$1.550'],
    ['KM252 - Disco Ceramica 7 1/4 Kamasa', '$7.057', '$1.800', '$12.990', '$11.000 - $14.990', '$12.990', '+$4.133', '+$4.133'],
    ['101074 - Disco Continuo 7 1/4 Kamasa', '$7.676', '$1.800', '$13.990', '$12.000 - $15.990', '$13.990', '+$4.514', '+$4.514'],
    ['DPG36 - Disco Silicio Pulir Grano 36', '$2.678', '$1.000', '$5.490', '$4.990 - $6.990', '$5.490', '+$1.812', '+$1.812'],
    ['PCM14236 - Piedra Cónica M14 #36', '$5.938', '$1.800', '$12.990', '$11.990 - $17.000', '$13.990', '+$5.252', '+$6.252 (+19%)'],
    ['DPAC115 - Disco Pulidor Inox 115mm', '$4.760', '$1.500', '$8.990', '$7.990 - $9.990', '$8.990', '+$2.730', '+$2.730'],
    ['106378 - Base Soporte Velcrada 4"', '$5.831', '$1.500', '$9.990', '$7.990 - $11.990', '$9.990', '+$2.659', '+$2.659'],
    ['101328 - Disco Lija / Flap 4.5"', '$1.821', '$600', '$3.800', '$2.990 - $4.500', '$3.490', '+$1.379', '+$1.069'],
    ['DTG115 - Disco Traslapado Mármol 4.5"', '$893', '$500', '$2.490', '$1.990 - $2.990', '$2.490', '+$1.097', '+$1.097'],
    ['104295 - Trompo Resina Diamantado', '$10.924', '$2.000', '$9.990', '$14.990 - $19.990', '$15.990', '-$2.934 (PÉRDIDA)', '+$3.066 (RENTABLE)'],
    ['SCDM35 - Sierra Copa Diamantada 35mm', '$9.223', '$2.000', '$17.000', '$12.990 - $18.990', '$16.990', '+$5.777', '+$5.767'],
    ['DR11410 - Disco Ranurador 114x10mm', '$38.080', '$4.000', '$59.990', '$49.990 - $69.990', '$54.990', '+$17.910', '+$12.910']
]

table_comp = doc.add_table(rows=len(data_comp) + 1, cols=len(headers_comp))
table_comp.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table_comp.rows[0].cells
for idx, title in enumerate(headers_comp):
    hdr_cells[idx].text = title
    set_cell_background(hdr_cells[idx], HEX_NAVY)
    set_cell_margins(hdr_cells[idx], top=80, bottom=80, left=50, right=50)
    for p in hdr_cells[idx].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(7.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

for row_idx, row_data in enumerate(data_comp):
    row_cells = table_comp.rows[row_idx + 1].cells
    bg_hex = HEX_ALT_ROW if row_idx % 2 == 1 else 'FFFFFF'
    for col_idx, cell_value in enumerate(row_data):
        row_cells[col_idx].text = cell_value
        set_cell_background(row_cells[col_idx], bg_hex)
        set_cell_margins(row_cells[col_idx], top=60, bottom=60, left=40, right=40)
        for p in row_cells[col_idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(7)
                r.font.color.rgb = COLOR_DARK
                if 'PÉRDIDA' in cell_value:
                    r.font.color.rgb = RGBColor(220, 38, 38)
                    r.font.bold = True
                elif 'RENTABLE' in cell_value:
                    r.font.color.rgb = RGBColor(34, 197, 94)
                    r.font.bold = True

doc.add_paragraph()

# Section 3: Key Strategic Findings
h3 = doc.add_heading(level=1)
r_h3 = h3.add_run('3. Hallazgos Críticos y Corrección de Precios Urgente')
r_h3.font.name = 'Arial'
r_h3.font.color.rgb = COLOR_NAVY
r_h3.font.bold = True

p3 = doc.add_paragraph()
r3 = p3.add_run(
    '1. Alerta Crítica en Trompo de Resina (Código `104295`):\n'
    '   - El costo de bodega con IVA es de **$10.924 CLP**. En la web estaba publicado a **$9.990 CLP**.\n'
    '   - Cada venta de este producto generaba una **pérdida neta de $2.934 CLP** (al incluir la publicidad de Meta Ads).\n'
    '   - En el mercado técnico (Mercado Libre/Marmolerías), este producto se vende entre **$14.990 y $19.990 CLP**. Recomendamos ajustarlo a **$15.990 CLP**, '
    'lo que convierte el producto en altamente rentable (+$3.066 CLP de ganancia neta por unidad).\n\n'
    '2. Oportunidad de Margen en Piedra Cónica M14 (`PCM14236`):\n'
    '   - Costo real con IVA: $5.938 CLP. Se vende en mercado hasta en $17.000 CLP. Recomendamos ajustar de $12.990 a $13.990 CLP para absorber publicidad sin perder conversión.\n\n'
    '3. Efecto de Venta en Paquete / Cotizador por Volumen:\n'
    '   - Cuando el cliente cotiza más de 10 o 25 unidades en el cotizador online, el gasto publicitario de Meta Ads ($1.500 aprox.) se diluye a centavos por unidad, '
    'permitiendo otorgar hasta un 25% de descuento por volumen manteniendo una alta rentabilidad para América Pulido.'
)
r3.font.name = 'Arial'
r3.font.size = Pt(10)
r3.font.color.rgb = COLOR_DARK

target_path = r'c:\Users\usuario\Desktop\america-pulido-web\Propuesta_Cotizador_y_Estrategia_Descuentos_America_Pulido.docx'
try:
    doc.save(target_path)
    print('Updated proposal DOCX with 15-item matrix.')
except Exception as e:
    alt_path = r'c:\Users\usuario\Desktop\america-pulido-web\Propuesta_Cotizador_y_Estrategia_Descuentos_America_Pulido_v2.docx'
    doc.save(alt_path)
    print('Saved to v2 due to lock.')
